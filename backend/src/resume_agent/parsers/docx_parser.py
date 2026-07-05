"""DOCX 文件解析器。

使用 ``python-docx`` 遍历段落与表格单元格提取文本。
对齐 design.md 第 2.2 节。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


def extract_text_from_docx(file_path: Path) -> str:
    """从 DOCX 文件提取全文文本。

    遍历文档所有段落（按出现顺序），段落间用换行符分隔。
    表格单元格内的段落文本也会被提取，保持与正文一致的拼接方式。

    Args:
        file_path: DOCX 文件路径。

    Returns:
        提取的纯文本。空文档返回空字符串。

    Raises:
        FileNotFoundError: 文件不存在。
        RuntimeError: python-docx 打开或解析失败。
    """
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX 文件不存在: {file_path}")

    try:
        document = Document(str(file_path))
    except Exception as exc:  # noqa: BLE001 - 转换为统一异常
        raise RuntimeError(f"打开 DOCX 失败: {file_path}: {exc}") from exc

    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if text:
                        parts.append(text)

    return "\n".join(parts)
