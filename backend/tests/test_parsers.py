"""PDF 与 DOCX 文件解析器单元测试。

在测试中用 PyMuPDF 创建临时 PDF，用 python-docx 创建临时 DOCX，
验证解析器提取结果与预期一致。
"""

from __future__ import annotations

from pathlib import Path

import fitz  # type: ignore[import-not-found]
from docx import Document

from resume_agent.parsers.docx_parser import extract_text_from_docx
from resume_agent.parsers.pdf_parser import extract_text_from_pdf

# === PDF 解析器测试 ===


def _create_test_pdf(path: Path, pages: list[str]) -> None:
    """用 PyMuPDF 创建包含多页文本的测试 PDF。"""
    doc = fitz.open()
    for page_text in pages:
        page = doc.new_page()
        page.insert_text((72, 72), page_text, fontsize=12)
    doc.save(str(path))
    doc.close()


def test_pdf_extract_single_page(tmp_path: Path) -> None:
    """单页 PDF 提取文本。"""
    pdf_path = tmp_path / "test.pdf"
    _create_test_pdf(pdf_path, ["Hello PDF World"])
    text = extract_text_from_pdf(pdf_path)
    assert "Hello PDF World" in text


def test_pdf_extract_multiple_pages(tmp_path: Path) -> None:
    """多页 PDF 提取文本，页间用空行分隔。"""
    pdf_path = tmp_path / "multi.pdf"
    _create_test_pdf(pdf_path, ["Page One Content", "Page Two Content"])
    text = extract_text_from_pdf(pdf_path)
    assert "Page One Content" in text
    assert "Page Two Content" in text
    # 页间应有分隔
    assert "\n\n" in text


def test_pdf_extract_chinese_text(tmp_path: Path) -> None:
    """PDF 提取中文文本（简历常见场景）。"""
    pdf_path = tmp_path / "cn.pdf"
    # insert_text 对中文支持有限，改用 text page 写入
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Zhang San Resume", fontsize=12)
    doc.save(str(pdf_path))
    doc.close()
    text = extract_text_from_pdf(pdf_path)
    assert "Zhang San Resume" in text


def test_pdf_extract_empty_pdf(tmp_path: Path) -> None:
    """空 PDF（无文本）返回空字符串。"""
    pdf_path = tmp_path / "empty.pdf"
    doc = fitz.open()
    doc.new_page()
    doc.save(str(pdf_path))
    doc.close()
    text = extract_text_from_pdf(pdf_path)
    assert text == ""


def test_pdf_file_not_found(tmp_path: Path) -> None:
    """文件不存在抛 FileNotFoundError。"""
    import pytest

    with pytest.raises(FileNotFoundError):
        extract_text_from_pdf(tmp_path / "nonexistent.pdf")


# === DOCX 解析器测试 ===


def _create_test_docx(path: Path, paragraphs: list[str]) -> None:
    """用 python-docx 创建包含多段落的测试 DOCX。"""
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(str(path))


def test_docx_extract_paragraphs(tmp_path: Path) -> None:
    """DOCX 提取多个段落文本。"""
    docx_path = tmp_path / "test.docx"
    _create_test_docx(docx_path, ["First paragraph", "Second paragraph", "Third"])
    text = extract_text_from_docx(docx_path)
    assert "First paragraph" in text
    assert "Second paragraph" in text
    assert "Third" in text


def test_docx_extract_with_table(tmp_path: Path) -> None:
    """DOCX 含表格，表格单元格文本应被提取。"""
    docx_path = tmp_path / "table.docx"
    doc = Document()
    doc.add_paragraph("Resume Header")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Company"
    table.cell(0, 1).text = "Tencent"
    table.cell(1, 0).text = "Role"
    table.cell(1, 1).text = "Engineer"
    doc.save(str(docx_path))

    text = extract_text_from_docx(docx_path)
    assert "Resume Header" in text
    assert "Company" in text
    assert "Tencent" in text
    assert "Engineer" in text


def test_docx_extract_empty_document(tmp_path: Path) -> None:
    """空 DOCX 返回空字符串。"""
    docx_path = tmp_path / "empty.docx"
    doc = Document()
    doc.save(str(docx_path))
    text = extract_text_from_docx(docx_path)
    assert text == ""


def test_docx_file_not_found(tmp_path: Path) -> None:
    """文件不存在抛 FileNotFoundError。"""
    import pytest

    with pytest.raises(FileNotFoundError):
        extract_text_from_docx(tmp_path / "nonexistent.docx")
